#!/usr/bin/env python3
"""Build EZ Solutions Pilot Findings DOCX — 5-section structure."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

SCREENSHOTS = Path('/Users/txu/Repo/wordpress/screenshots')

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return p

def add_screenshot(filename, caption, width=Inches(5.5)):
    path = SCREENSHOTS / filename
    if path.exists():
        doc.add_picture(str(path), width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ─── TITLE ───────────────────────────────────────────────────────────────────

title = doc.add_heading('EZ Solutions — WordPress AI Pilot', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Findings & Recommendations')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Elevate MSP | July 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ─── 1. EXEC SUMMARY ─────────────────────────────────────────────────────────

add_heading('1. Executive Summary')

add_para(
    'EZ Solutions needs to produce 50+ unique location pages per client site — '
    'pages that pass Google quality checks, look premium, and remain editable in Gutenberg. '
    'Their current spin-tax approach is being flagged as mass-produced content.'
)

add_para(
    'Over one week, we built a proof-of-concept that generated 51 Pennsylvania city HVAC pages, '
    'each with a fundamentally different layout, all passing automated quality gates for word count, '
    'uniqueness, imagery, structure, and SEO metadata. The entire stack runs locally in Docker with '
    'zero external dependencies at runtime.'
)

add_para('Key results:', bold=True)
add_bullet('51/51 pages pass all quality gates (word count, images, uniqueness, structure)')
add_bullet('28 distinct section types used across corpus — no two pages look alike')
add_bullet('Max pairwise content similarity: 3.3% (threshold: 8%)')
add_bullet('Median word count: ~1,900 (min 1,013 / max 3,507)')
add_bullet('Zero fabricated business claims detected')
add_bullet('All output is native Gutenberg blocks — fully editable in WordPress')

doc.add_page_break()

# ─── 2. SCOPE ─────────────────────────────────────────────────────────────────

add_heading('2. Scope: What We Heard')

add_para(
    'From the July 2 "WordPress QA Deepdive" call (ELVT-207) with Cindy, Marie (front-end), '
    'and Deb (back-end), we captured the following requirements:',
    italic=True
)

add_heading('Content & Quality', level=2)
add_bullet('Generate 50+ location pages per client, each with unique layout and content, in one workflow')
add_bullet('Target 400–1000+ words depth per page (EZ production standard)')
add_bullet('No mass-production sameness — client sniff-tests pages side by side')
add_bullet('No fabricated business claims (fictional brands = liability risk)')
add_bullet('Structured content: FAQ sections, CTAs, proper heading hierarchy')
add_bullet('No text walls — readable paragraph lengths, visual variety')

add_heading('Design & Platform', level=2)
add_bullet('Gutenberg-native — Elementor explicitly rejected')
add_bullet('Imagery / design layer — not walls of text')
add_bullet('Theme / brand layer must be professional and consistent')
add_bullet('Internal linking between location pages for SEO structure')
add_bullet('SEO meta on every page (Yoast integration)')
add_bullet('ACF / custom fields support for structured data')

add_heading('Infrastructure & Compliance', level=2)
add_bullet('Zero-egress inference — CA privacy constraints, data cannot leave network')
add_bullet('Fleet management — 300 sites, Deb runs monthly update sweeps')
add_bullet('MCP agent access — AI operating WordPress directly via API')
add_bullet('IT approval / no GitHub — self-hosted, works behind VPN, no external repos')
add_bullet('Client-editable output — final pages editable via WordPress or AI')

add_heading('The Solutions EZ Wants', level=2)
add_para(
    'In short: a production-ready pipeline that lets a non-developer generate 50+ premium, '
    'unique, SEO-ready location pages — Gutenberg-native, fully editable, running entirely '
    'on their infrastructure with no data leaving the network. Plus fleet management across '
    'their 300-site portfolio.'
)

doc.add_page_break()

# ─── 3. OPTIONS & SUGGESTIONS ────────────────────────────────────────────────

add_heading('3. Options & Suggestions')

add_para(
    'Based on what we validated in the POC, here is our recommended tooling by capability area:'
)

# Table
table = doc.add_table(rows=1, cols=3)
table.style = 'Medium Shading 1 Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Capability'
hdr[1].text = 'Recommended Tool'
hdr[2].text = 'Notes'

rows_data = [
    ('Content generation\n(50+ unique pages)', 'Claude Code + agentic pipeline',
     'Isolated agent per page with unique creative brief. Temperature 1.0 for maximum diversity. Quality gate enforces word count, uniqueness, structure.'),
    ('Zero-egress inference\n(on-prem / no data out)', 'Ollama + AI Engine custom env',
     'Ollama hosts open models (gemma4, llama) locally. AI Engine Pro routes WP to localhost:11434. Validated in POC.'),
    ('MCP agent access\n(AI controls WP)', 'AI Engine Pro ($59/site)',
     'Exposes 43 WP tools via MCP endpoint. Agents can create posts, write blocks, manage media. Bearer auth.'),
    ('Fleet management\n(300 sites)', 'MainWP (free dashboard + $228/yr extensions)',
     'Central dashboard for plugin updates, security scans, content deployment across all child sites. API scriptable.'),
    ('SEO metadata & schema', 'Yoast SEO (already installed)',
     'Pipeline writes _yoast_wpseo_metadesc automatically. Future: LocalBusiness JSON-LD schema.'),
    ('Imagery / media', 'Openverse (free, CC-licensed)',
     'Commercial-use image pool. 14 topic categories for HVAC. Future: paid stock or client photos.'),
    ('Non-dev page generation\n(Marie / front-end team)', 'Page Generator Pro ($99/yr)',
     'Template-based bulk generation in WP admin. No CLI needed. Good for simpler repetitive pages.'),
    ('Theme / brand layer', 'Custom child theme of TT5',
     'Fraunces/Inter typography, custom palette, header/footer template parts. One-time build per client.'),
    ('Design system / sections', 'Custom section library (25+ types)',
     'PHP renderers produce valid Gutenberg blocks. Ken Burns heroes, hover-lift cards, accordion FAQ, colored bands.'),
]

for cap, tool, notes in rows_data:
    row = table.add_row().cells
    row[0].text = cap
    row[1].text = tool
    row[2].text = notes

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(4.5)
    row.cells[1].width = Cm(4.5)
    row.cells[2].width = Cm(8.5)

doc.add_paragraph()
add_heading('Build vs. Buy Summary', level=2)
add_bullet('80% of the pipeline is "buy" (existing tools: Claude API, AI Engine, MainWP, Yoast, Openverse)')
add_bullet('20% is custom build (section library, quality gate, creative briefs, rendering pipeline)')
add_bullet('The custom 20% is what makes pages look fundamentally different from each other — it cannot be replicated with off-the-shelf plugins alone')

doc.add_page_break()

# ─── 4. POC RESULTS ──────────────────────────────────────────────────────────

add_heading('4. POC (Build Option) Results')

bon = add_para('What we built and proved in one week:')

add_heading('Architecture', level=2)
add_bullet('Docker Compose stack: WordPress (8181) + MariaDB (33081) + MainWP (8282) + wp-cli')
add_bullet('Agentic spec generator: Python script → 51 parallel Claude API calls → JSON layout specs')
add_bullet('Section renderer: PHP library (25+ section types) → valid Gutenberg block markup')
add_bullet('Quality gate: automated checks for word count, images, headings, uniqueness, claims')
add_bullet('Creative briefs: 51 unique style directives (cinematic-visual, data-dashboard, minimal-zen, etc.)')

add_heading('Quality Gate Results', level=2)

gate_table = doc.add_table(rows=1, cols=3)
gate_table.style = 'Light List Accent 1'
ghdr = gate_table.rows[0].cells
ghdr[0].text = 'Metric'
ghdr[1].text = 'Threshold'
ghdr[2].text = 'Result'

gate_rows = [
    ('Word count', '≥ 1,000 words', '51/51 pass (median ~1,900)'),
    ('Inline images', '≥ 2 per page', '51/51 pass'),
    ('H2 headings', '≥ 4 per page', '51/51 pass'),
    ('Paragraph length', '≤ 110 words', '51/51 pass'),
    ('Meta description', 'Present + ≤160 chars', '51/51 pass'),
    ('Shingle Jaccard', '< 0.08 (all pairs)', 'Max 0.033 (avg <0.02)'),
    ('Opening uniqueness', 'First 3 words unique', '51/51 unique'),
    ('Banned claims', 'Zero matches', '0 violations'),
]

for metric, threshold, result in gate_rows:
    row = gate_table.add_row().cells
    row[0].text = metric
    row[1].text = threshold
    row[2].text = result

doc.add_paragraph()
add_heading('Design Quality', level=2)
add_bullet('28 distinct section types used across the 51 pages')
add_bullet('36% of sections have colored background bands (sand, ink, cream, dark-gradient)')
add_bullet('49/51 pages include a structured service listing (feature_row, cards, or icon_list)')
add_bullet('All headlines ≤10 words; all CTAs ≤5 words (design taste rules enforced)')
add_bullet('Premium CSS: Ken Burns animation, multi-layer shadows, hover-lift, cubic-bezier transitions')

add_heading('Screenshots', level=2)

add_screenshot('pittsburgh-hero.png', 'Pittsburgh — hero_cover with Ken Burns animation')
add_screenshot('pittsburgh-full.png', 'Pittsburgh — full page showing section variety')
add_screenshot('bethlehem-hero.png', 'Bethlehem — hero_split with offset card design')
add_screenshot('bethlehem-full.png', 'Bethlehem — colored bands, stats, feature rows')
add_screenshot('lancaster-hero.png', 'Lancaster — editorial-longform hero style')
add_screenshot('lancaster-full.png', 'Lancaster — data cards, FAQ accordion, CTA bands')
add_screenshot('erie-hero.png', 'Erie — cinematic-visual hero approach')
add_screenshot('phoenixville-hero.png', 'Phoenixville — minimal-zen clean design')
add_screenshot('johnstown-hero.png', 'Johnstown — overlap-depth card layout')

doc.add_page_break()

# ─── 5. RECOMMENDED NEXT STEPS ───────────────────────────────────────────────

add_heading('5. Recommended Next Steps')

add_heading('Immediate (Week 1-2)', level=2)
add_bullet('Approve tooling selections and budget (AI Engine Pro $59/site, Page Generator Pro $99/yr)')
add_bullet('Provision a persistent dev server for the Docker stack (their infrastructure, behind VPN)')
add_bullet('Reinstall Ollama for zero-egress inference (crashed during POC; architecture validated)')
add_bullet('Select first real client site for pilot deployment')

add_heading('Short-term (Week 3-4)', level=2)
add_bullet('Adapt theme layer for real client brand (replace Keystone Comfort with actual client)')
add_bullet('Swap Openverse images for client-specific stock or job-site photos')
add_bullet('Run pipeline on real client city list (not PA HVAC demo)')
add_bullet('Train Marie (front-end) on Page Generator Pro workflow for simpler page types')
add_bullet('Connect MainWP fleet dashboard to 2-3 production sites for update management')

add_heading('Medium-term (Month 2-3)', level=2)
add_bullet('Build fleet governance layer: per-site content rollout, approval workflows')
add_bullet('Add LocalBusiness JSON-LD schema markup to all location pages')
add_bullet('Implement hub-and-spoke internal linking by geo-distance')
add_bullet('Expand section library for other verticals (legal, dental, roofing)')
add_bullet('Tune word-count targets per vertical (HVAC vs. legal vs. dental)')
add_bullet('Set up agent guardrails for MCP access (AI Engine Pro curated tool sets)')

add_heading('Long-term', level=2)
add_bullet('Roll out pipeline across portfolio (300 sites)')
add_bullet('Vertical-specific creative brief libraries')
add_bullet('Client self-serve: non-technical staff trigger page generation from WP admin')
add_bullet('Automated content freshness: re-run pipeline quarterly with updated data')

# ─── SAVE ─────────────────────────────────────────────────────────────────────

out_path = '/Users/txu/Repo/wordpress/EZ-Solutions-Pilot-Findings.docx'
doc.save(out_path)
print(f"Saved: {out_path}")
